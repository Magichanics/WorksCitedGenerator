'''
Author: Jan Garong
Date: February 6th, 2019

format:
https://owl.purdue.edu/owl/research_and_citation/apa_style/apa_formatting_and_style_guide/general_format.html

Some notes:
Works Cited (centered)
Margin is 1
Author, "Title" Publisher, (Year, Month Day), URL

please use test_history_df.fillna(0) for nulls
'''

from docx import Document # docx is python-docx
from docx.shared import Pt
from docx.text.parfmt import ParagraphFormat
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import numpy as np
import datetime

class WorksCited:

    # sets font to Times New Roman, 12
    def set_font(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    # add 0.5 in. margins
    def set_margins(self):
        sections = self.document.sections
        for s in sections:
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)

    # set format (double spacing % hanging indentations)
    def set_format(self, is_title=False):
        p_format = self.document.styles['Normal'].paragraph_format
        p_format.line_spacing_rule = 2
        p_format.first_line_indent = Inches(-0.25)

    def time_to_string(self, datetime_input):

        # return datetime_input.strftime("(%Y, %B " + str(int("%d")) + ")")

        # convert to list
        date_lst = datetime_input.strftime("%Y-%B-%d").split('-')

        # get rid of trailing zeros
        date_lst[2] = str(int(date_lst[2]))

        # combine them together (use regex please)
        return '(' + date_lst[0] + ', ' + date_lst[1] + ' ' + date_lst[2] + ')'

    def get_author_and_date(self, authors, date):

        # check if there are no known authors
        if authors == 0:
            authors = ''  # leave it blank
        else:  # add trailing space between date and name
            authors += ' '

        # see if you can get the date and time according to APA
        try:
            date_string = self.time_to_string(date)
        except:
            date_string = '(n.d)'

        return authors, date_string

    # similar to add_citation(), but returns a string value. Use this when doing alphabetical sorting.
    def add_mock_citation(self, authors, date, name, url):

        # check if there are no known authors
        authors, date_string = self.get_author_and_date(authors, date)

        # return string citation
        return authors + date_string + '. ' + str(name) + '.' + ' Retrieved from ' + url

    # add a new citation (needs name, date of publication/date of access, url, author/publisher)
    def add_citation(self, authors, date, name, url):

        # check if there are no known authors
        authors, date_string = self.get_author_and_date(authors, date)

        # I don't think there will be an instance where the title or url is going to be missing, but just in case, leave
        # the if statements for them right here.

        # create citation
        p = self.document.add_paragraph(authors + date_string + '. ')
        p.add_run(str(name) + '.').italic = True
        p.add_run(' Retrieved from ' + url)

    # save document
    def save(self, save_location='WorksCited.docx'):
        self.document.save(save_location)

    # initialize works cited page
    def __init__(self):

        # create title page
        self.document = Document()

        p = self.document.add_paragraph("Works Cited")

        # use double spacing, and center page
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_format()
        self.set_font()

        # set margins
        self.set_margins()


