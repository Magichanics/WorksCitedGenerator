'''
Name: Magichanics
Date: February 6th, 2019

format:
https://owl.purdue.edu/owl/research_and_citation/apa_style/apa_formatting_and_style_guide/general_format.html

Some notes:
Works Cited (centered)
Margin is 0.5
Author, "Title" Publisher, (Year, Month Day), URL

do 0.5 inch from all sides.

please use test_history_df.fillna(0) for nulls
'''

from docx import Document
from docx.shared import Pt
from docx.text.parfmt import ParagraphFormat
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import numpy as np
import datetime

class WorksCited:

    # month integers to month names
    to_month = {'01': 'January',
                '02': 'February',
                '03': 'March',
                '04': 'April',
                '05': 'May',
                '06': 'June',
                '07': 'July',
                '08': 'August',
                '09': 'September',
                '10': 'October',
                '11': 'November',
                '12': 'December'}

    # sets font to Times New Roman, 12
    def set_font(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    # add 0.5 in. margins
    def set_margins(self):
        sections = self.document.sections
        for s in sections:
            s.top_margin = Inches(0.5)
            s.bottom_margin = Inches(0.5)
            s.left_margin = Inches(0.5)
            s.right_margin = Inches(0.5)

    def time_to_string(self, datetime_input):

        # convert to list
        date_string = datetime_input.strftime("%Y-%b-%d")
        date_lst = date_string.split('-')

        # use dictionary to get name
        date_lst[1] = to_month[date_lst[1]]

        # get rid of trailing zeros
        date_lst[2] = str(int(date_lst[2]))

        # combine them together (use regex please)
        return '(' + date_lst[0] + ', ' + date_lst[1] + ' ' + date_lst[2] + ')'

    # add a new citation (needs name, date of publication/date of access, url, author/publisher)
    def add_citation(self, authors, date, name, url):

        # check if there are no known authors
        if authors == 0:
            authors = '' # lave it blank
        else: # add trailing space between date and name
            authors += ' '

        # check if there is a date
        if date == 0:
            date_string = '(n.d)'
        else: # convert datetime to string, formatted accordingly to APA
            date_string = self.time_to_string(date)

        # I don't think there will be an instance where the title or url is going to be missing, but just in case, leave
        # the if statements for them right here.

        # create citation
        p = self.document.add_paragraph(authors + date_string + '. ')
        p.add_run(str(name) + '.').italic = True
        p.add_run(' Retrieved from ' + url)

    # save document
    def save(self):
        self.document.save('WorksCited.docx')

    # initialize works cited page
    def __init__(self):

        # create title page
        self.document = Document()

        # set font
        self.set_font()
        p = self.document.add_paragraph("Works Cited")
        p.style = self.document.styles['Normal']

        # use double spacing, and center page
        p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # set margins
        self.set_margins()

